import webbrowser
from docx import Document

# TODO: figure out venv implementation so that installing packages aren't required

"""
----------------------------------------------------------------------------
Globals
----------------------------------------------------------------------------
"""

# global to store recipes
new_data = {"hits": []}
selected_data = {"hits": []}


def create_recipe_document():
    total_recipes = []
    just_ingredients = []
    modified_data = [total_recipes, just_ingredients]

    # create new document
    doc = Document()
    # add heading
    doc.add_heading("Saved Recipes")

    for i, recipe_data in enumerate(new_data['hits'], start=1):
        # rename the recipe so they go in ascending order
        new_name = 'recipe ' + str(i)

        # create the new recipe dictionary
        new_recipe = {new_name: recipe_data}

        # add new recipe to list
        total_recipes.append(new_recipe)

        # isolate the ingredients list from the recipe
        new_ingredients = {new_name + ' ingredients': recipe_data['recipe']['ingredientLines']}

        # add isolated ingredients to list of ingredients
        just_ingredients.append(new_ingredients)

    # add each recipe to document
    for recipe_info in modified_data[0]:
        for result_number, recipe_details in recipe_info.items():
            title_paragraph = doc.add_paragraph()
            runner = title_paragraph.add_run(f"Recipe Title: {recipe_details['recipe']['label']}")
            runner.bold = True
            doc.add_paragraph(f"URL: {recipe_details['recipe']['url']}")

            # add ingredients as bulleted list
            doc.add_paragraph(f"Ingredients: ")
            for each_ingredient in recipe_details['recipe']['ingredientLines']:
                ingredient_paragraph = doc.add_paragraph(f"{each_ingredient}")
                ingredient_paragraph.style = 'List Bullet'

            doc.add_paragraph("\n")

    response_filename = 'Recipes.docx'
    doc.save(response_filename)


def create_ingredients_document():
    total_recipes = []
    just_ingredients = []
    modified_data = [total_recipes, just_ingredients]
    for i, recipe_data in enumerate(new_data['hits'], start=1):
        # rename the recipe so they go in ascending order
        new_name = 'recipe ' + str(i)

        # create the new recipe dictionary
        new_recipe = {new_name: recipe_data}

        # add new recipe to list
        total_recipes.append(new_recipe)

        # isolate the ingredients list from the recipe
        new_ingredients = {new_name + ' ingredients': recipe_data['recipe']['ingredientLines']}

        # add isolated ingredients to list of ingredients
        just_ingredients.append(new_ingredients)

    # create new document
    doc = Document()
    # add heading
    doc.add_heading("Ingredients List")

    # add each recipe to document
    for recipe_info in modified_data[0]:
        for result_number, recipe_details in recipe_info.items():
            for recipe_ingredient in recipe_details['recipe']['ingredientLines']:
                doc.add_paragraph(f"{recipe_ingredient}")

    response_filename = 'Ingredients List.docx'
    doc.save(response_filename)


def open_url(url):
    import webbrowser
    webbrowser.open(url)


def argument_handler(*args):
    if not args:
        return "You need to provide at least 1 argument."
    else:
        return list(args)


def remove_str_chars(input_string, num):
    if num >= 0:
        return input_string[:-num]
    else:
        return input_string


def process_food_list(input_list):
    request_string = ""
    for food in input_list:
        request_string = request_string + food + "%2c%20"
    # remove the last %2c%20 from the collated string which is unicode for ", "
    api_string = remove_str_chars(request_string, 6)
    return api_string


def browse_recipes():
    url = "https://www.allrecipes.com/"
    webbrowser.open(url)
