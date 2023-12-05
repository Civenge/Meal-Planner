import json
import requests
import random
from art import *
import webbrowser
from send_and_receive_class import TcpConnect
from docx import Document

"""
This project uses the edamam API to obtain recipes, with the API information below:
    api here: https://developer.edamam.com//admin/applications/1409623863337
    api doc: https://developer.edamam.com/edamam-docs-recipe-api
    
Requires installation of art via "pip install art" in console.
Requires installation of docx via "pip install python-docx" in console.
"""
# -----------------------------------------------------------------------------------------------------------------
# SETUP TCP CONNECTION WITH TCPCONNECT CLASS
server_name = 'localhost'
server_port = 1088

my_session = TcpConnect('User', server_name, server_port)
my_session.connect()

# -----------------------------------------------------------------------------------------------------------------

tprint("Meal Planner")
# unicode is for italics and green
print("\033[3m\033[92mThis program will take approximately 30 seconds to complete for a single recipe.\033[0m")
print("* Use this program to pick out recipes and save a list of new recipes to try!")
print("* Start by inputting some ingredients that you want to use to find a recipe.")
print("* There will be an opportunity to export any saved recipes at the end of the search.")
print("* Would you like more information about this program prior to starting the search? (Yes or No)\n")

# global to store recipes
new_data = {"hits": []}


def argument_handler(*args):
    if not args:
        return "You need to provide at least 1 argument."
    else:
        return list(args)


def remove_str_chars(input_string, num):
    if num >= 0:
        return input_string[:-num]
    else:
        return input_string


def process_food_list(input_list):
    request_string = ""
    for food in input_list:
        request_string = request_string + food + "%2c%20"
    # remove the last %2c%20 from the collated string which is unicode for ", "
    api_string = remove_str_chars(request_string, 6)
    return api_string


def create_recipe_document(recipe_list):
    # create new document
    doc = Document()
    # add heading
    doc.add_heading("Saved Recipes")

    # add each recipe to document
    for recipe_info in recipe_list[0]:
        for result_number, recipe_details in recipe_info.items():
            title_paragraph = doc.add_paragraph()
            runner = title_paragraph.add_run(f"Recipe Title: {recipe_details['recipe']['label']}")
            runner.bold = True
            doc.add_paragraph(f"URL: {recipe_details['recipe']['url']}")

            # add ingredients as bulleted list
            doc.add_paragraph(f"Ingredients: ")
            for each_ingredient in recipe_details['recipe']['ingredientLines']:
                ingredient_paragraph = doc.add_paragraph(f"{each_ingredient}")
                ingredient_paragraph.style = 'List Bullet'

            doc.add_paragraph("\n")

    response_filename = 'Recipes.docx'
    doc.save(response_filename)


def create_ingredients_document(formatted_data):
    # create new document
    doc = Document()
    # add heading
    doc.add_heading("Ingredients List")

    # add each recipe to document
    for recipe_info in formatted_data[0]:
        for result_number, recipe_details in recipe_info.items():
            for recipe_ingredient in recipe_details['recipe']['ingredientLines']:
                doc.add_paragraph(f"{recipe_ingredient}")

    response_filename = 'Ingredients List.docx'
    doc.save(response_filename)


user_input = input().lower()
if user_input == "yes" or user_input == "y":
    print("This is a program that will take in 5 inputs from a user to help you find a recipe.")
    print("The first input will be a list of ingredients that you want to include, separated by commas.")
    print("Then the program will ask for how many recipes you want to see, between 1 and 20.")
    print("The next input will be an ingredient that you want to exclude, or you can leave it blank.")
    print("Then you can choose which recipe results you want to save based on the displayed recipe number.")
    print("Finally, you will have the option to do another search and save even more recipes, or get the results")
    print("  to download as a word document.")
    print("Let's get started!\n")

browse_recipes = input(
    'Would you like to search for recipes or browse for recipes? Type \'search\' or \'browse\'\n').lower()
if browse_recipes.lower() == "browse" or browse_recipes.lower() == "b":
    url = "https://www.allrecipes.com/"
    webbrowser.open(url)

    # sends dummy JSON so microservice will end elegantly
    reply = my_session.send_data(new_data)
    my_session.end_send()

    # obtain data from microservice
    modified_data = my_session.get_mod_data()

    # finally, close the TCP connection
    my_session.disconnect()
    print("Enjoy your browsing, goodbye!")
    exit(0)

# create variables to store ingredients and saved recipes between loop runs
ingredients = None
selected_data = None


while True:
    food_list = None
    while food_list is None:
        ingredients = input("What ingredients do you want to use (chicken, cheese, etc)?\n")
        if ingredients:
            food_list = argument_handler(ingredients)
            formatted_string = process_food_list(food_list)
            test_response = requests.get(
                "https://api.edamam.com/api/recipes/v2?type=public&q=" + formatted_string +
                "&app_id=2286dd85&app_key=1cdfcd395ccf99e349b18f54eaa4416f&random=true&field=url&field=label"
                "&field=ingredientLines")
            dict_from_json = json.loads(test_response.text)
            if not dict_from_json["hits"]:
                print(f"\033[1m\033[91mYour search for {ingredients} found no recipes, please try again.\033[0m")
                food_list = None
        else:
            print("Please select at least one ingredient.\n")

    num_recipes = None
    while num_recipes is None:
        recipe_count = input("How many recipes would you like to view? (1 to 20)\n")
        if recipe_count.isdigit():
            num_recipes = int(recipe_count)
            if 1 <= num_recipes <= 20:
                break
        print(f"\033[1m\033[91mPlease pick a number between 1 and 20.\033[0m\n")
        recipe_count = ""
        num_recipes = None

    formatted_string = process_food_list(food_list)

    excluded_ingredients_str = input(
        "Please type out any ingredients you want excluded from the list (rice, bread), or press Enter to skip.\n")
    excluded_ingredients_list = argument_handler(excluded_ingredients_str)

    if excluded_ingredients_list:
        print(f"Excluding these ingredients:{excluded_ingredients_list}")

    else:
        excluded_ingredients_str = ""

    # build string for api call
    response = requests.get("https://api.edamam.com/api/recipes/v2?type=public&q=" + formatted_string +
                            "&app_id=2286dd85&app_key=1cdfcd395ccf99e349b18f54eaa4416f&" + excluded_ingredients_str +
                            "&random=true&field=url&field=label&field=ingredientLines")

    if response.status_code == 200:
        # parse the json
        dict_from_json = json.loads(response.text)
        if not dict_from_json["hits"]:
            print(f"\033[1m\033[91mYour search for {ingredients} found no recipes, please try again.\033[0m")
            exit(1)
        selected_recipes = random.sample(dict_from_json["hits"], num_recipes)
        selected_data = {
             "hits": selected_recipes
        }

        # prints out the nicely formatted recipe results
        for i, recipe_data in enumerate(selected_data["hits"], start=1):
            recipe = recipe_data["recipe"]
            recipe_url = recipe["url"]
            recipe_name = recipe["label"]
            ingredients = recipe_data["recipe"]["ingredientLines"]

            print(f"Recipe{i}: {recipe_name}")
            print(f"Url: {recipe_url}")
            for ingredient in ingredients:
                print(f"  {ingredient}")

    else:
        print("API request failed with status code: ", response.status_code)

    ask_add_to_shopping_list = input("Would you like to add any of these recipes to your saved recipes? (Yes or No)\n")
    if ask_add_to_shopping_list.lower() == "y" or ask_add_to_shopping_list.lower() == "yes":

        # loop to cover saving recipes
        while True:
            ask_add_which_recipes = input("What recipes would you like to save? (Ex: 1, 2, 4)\n")

            # remove commas from input
            split_values = ask_add_which_recipes.split(',')

            # remove whitespace in list
            current_saved_recipe_list = []
            for x in split_values:
                stripped_x = x.strip()
                current_saved_recipe_list.append(stripped_x)

            # removes non number values, but leaves a single negative sign if found (-)
            integer_list = []
            for x in current_saved_recipe_list:
                stripped_x = x.lstrip('-')
                if stripped_x.isdigit():
                    integer_list.append(int(x))

            # check for negative inputs
            if int(min(integer_list)) <= 0:
                print(
                    f"\033[1m\033[91mPlease select more than 0 recipes and make sure every selection is valid.\033[0m")
                continue

            # verifies save choices are within range and selection isn't greater than recipe count
            if len(integer_list) <= num_recipes and int(max(integer_list)) <= num_recipes:
                print(f"Here is what you selected: {integer_list}\n")
                print("******* Adding recipes to saved recipes... *******\n")

                for idx in range(len(integer_list)):
                    new_data["hits"].append(selected_data["hits"][integer_list[idx] - 1])
                break

            else:
                print(f"\033[1m\033[91mPlease select less recipes or make sure every selection is valid.\033[0m")

    ask_another_recipe = input("Would you like to search for another recipe? (Yes or No)\n")
    if ask_another_recipe.lower() == "y" or ask_another_recipe.lower() == "yes":
        continue

    else:
        print("******* Processing Request, Please Wait... *******\n")
        reply = my_session.send_data(new_data)
        print("Received reply from microservice: ", reply, '\r\n')
        my_session.end_send()

        # obtain data from microservice
        modified_data = my_session.get_mod_data()

        # close socket and print results
        print('Received all recipes list:', '\r\n', modified_data[0], '\r\n')
        print('Received all ingredients list:', '\r\n', modified_data[1])

        # finally, close the TCP connection
        my_session.disconnect()

        # prompt for storing recipes
        ask_store_recipes = input("Would you like the saved recipes in a Word Document? (Yes or No)\n")
        if ask_store_recipes.lower() == "yes" or ask_store_recipes.lower() == "y":

            storing_recipes = True
            while storing_recipes is True:

                # prompt for what to store
                print("Enter 1 for recipes only.")
                print("Enter 2 for an ingredients list only.")
                print("Enter 3 for both.")
                print("Enter 0 to exit.")
                ask_what_store = input("What would you like to store?\n")
                if ask_what_store == "1":
                    print("Creating \"Recipes.docx\"")
                    # function to print out the saved recipe in a Word document
                    create_recipe_document(modified_data)
                    break

                elif ask_what_store == "2":
                    print("Creating \"Ingredients List.docx\"")
                    # function to print out the ingredients list in a Word document
                    create_ingredients_document(modified_data)
                    break

                elif ask_what_store == "3":
                    print("Creating both \"Recipes.docx\" and \"Ingredients List.docx\"")
                    create_recipe_document(modified_data)
                    create_ingredients_document(modified_data)
                    break

                elif ask_what_store == "0":
                    break

                else:
                    print(f"\033[1m\033[91mPlease enter 1, 2, 3, or type 0 to exit.\033[0m")
                    continue

        print("Goodbye!")
        break
