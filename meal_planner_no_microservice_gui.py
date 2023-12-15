import json
import requests
import random
from tkinter import *
from tkinter import scrolledtext, Tk, Button, messagebox
import webbrowser
from docx import Document

"""
This project uses the edamam API to obtain recipes, with the API information below:
    api here: https://developer.edamam.com//admin/applications/1409623863337
    api doc: https://developer.edamam.com/edamam-docs-recipe-api

Requires installation of docx via "pip install python-docx" in console.
"""
"""
----------------------------------------------------------------------------
Globals
----------------------------------------------------------------------------
"""
# global to store recipes
new_data = {"hits": []}
selected_data = {"hits": []}

"""
----------------------------------------------------------------------------
Functions
----------------------------------------------------------------------------
"""


def argument_handler(*args):
    if not args:
        return "You need to provide at least 1 argument."
    else:
        return list(args)


def remove_str_chars(input_string, num):
    if num >= 0:
        return input_string[:-num]
    else:
        return input_string


def process_food_list(input_list):
    request_string = ""
    for food in input_list:
        request_string = request_string + food + "%2c%20"
    # remove the last %2c%20 from the collated string which is unicode for ", "
    api_string = remove_str_chars(request_string, 6)
    return api_string


def create_ingredients_document(formatted_data):
    # create new document
    doc = Document()
    # add heading
    doc.add_heading("Ingredients List")

    # add each recipe to document
    for recipe_info in formatted_data[0]:
        for result_number, recipe_details in recipe_info.items():
            for recipe_ingredient in recipe_details['recipe']['ingredientLines']:
                doc.add_paragraph(f"{recipe_ingredient}")

    response_filename = 'Ingredients List.docx'
    doc.save(response_filename)


def browse_recipes():
    url = "https://www.allrecipes.com/"
    webbrowser.open(url)


def search_recipes():
    global selected_data
    # clear existing text
    output_text.delete(1.0, END)
    excluded_ingredients_str = excluded_ingredients_entry.get()
    ingredients = ingredients_entry.get()
    if not ingredients:
        output_text.insert(END, "Please enter the number of recipes.\n")
        return
    num_recipes = int(num_recipes_entry.get())
    # Add your logic for searching recipes here
    if ingredients:
        food_list = argument_handler(ingredients)
        formatted_string = process_food_list(food_list)
        test_response = requests.get(
            "https://api.edamam.com/api/recipes/v2?type=public&q=" + formatted_string +
            "&app_id=2286dd85&app_key=1cdfcd395ccf99e349b18f54eaa4416f&random=true&field=url&field=label"
            "&field=ingredientLines")
        dict_from_json = json.loads(test_response.text)
        if not dict_from_json["hits"]:
            output_text.insert(END, f"Your search for {ingredients} found no recipes, please try again.")
            return
    else:
        output_text.insert(END, "Please select at least one ingredient.\n")
        return
    formatted_string = process_food_list(food_list)

    response = requests.get("https://api.edamam.com/api/recipes/v2?type=public&q=" + formatted_string +
                            "&app_id=2286dd85&app_key=1cdfcd395ccf99e349b18f54eaa4416f&" + excluded_ingredients_str +
                            "&random=true&field=url&field=label&field=ingredientLines")
    if response.status_code == 200:
        # parse the json
        dict_from_json = json.loads(response.text)
        if not dict_from_json["hits"]:
            output_text.insert(END, f"Your search for {ingredients} found no recipes, please try again.")
            exit(1)
        selected_recipes = random.sample(dict_from_json["hits"], num_recipes)
        selected_data = {
             "hits": selected_recipes
        }
        for i, recipe_data in enumerate(selected_data["hits"], start=1):
            recipe = recipe_data["recipe"]
            recipe_url = recipe["url"]
            recipe_name = recipe["label"]
            ingredients = recipe_data["recipe"]["ingredientLines"]

            # Append the recipe information to the text widget
            output_text.insert(END, f"Recipe{i}: {recipe_name}\n")
            output_text.insert(END, f"Url: {recipe_url}\n")

            for ingredient in ingredients:
                output_text.insert(END, f"  {ingredient}\n")

            output_text.insert(END, "\n")

    else:
        output_text.insert(END, f"API request failed with status code: {response.status_code}")


def save_recipes():
    global selected_data
    global new_data
    global output_text
    saved_recipes = saved_recipes_entry.get()
    if not saved_recipes:
        output_text.insert(END, f"Please select more than 0 recipes and make sure every selection is valid.")
        return
    split_values = saved_recipes.split(',')
    # remove whitespace in list
    current_saved_recipe_list = []
    for x in split_values:
        stripped_x = x.strip()
        current_saved_recipe_list.append(stripped_x)

    # removes non number values, but leaves a single negative sign if found (-)
    integer_list = []
    for x in current_saved_recipe_list:
        stripped_x = x.lstrip('-')
        if stripped_x.isdigit():
            integer_list.append(int(x))

    if int(min(integer_list)) <= 0:
        output_text.insert(END, f"Please select more than 0 recipes and make sure every selection is valid.")
        return

    output_text.insert(END, f"Here is what you selected: {integer_list}\n")
    output_text.insert(END, "******* Adding recipes to saved recipes... *******\n")

    for idx in range(len(integer_list)):
        new_data["hits"].append(selected_data["hits"][integer_list[idx] - 1])

    return


def create_recipe_document():
    total_recipes = []
    just_ingredients = []
    modified_data = [total_recipes, just_ingredients]
    for i, recipe_data in enumerate(new_data['hits'], start=1):
        # rename the recipe so they go in ascending order
        new_name = 'recipe ' + str(i)

        # create the new recipe dictionary
        new_recipe = {new_name: recipe_data}

        # add new recipe to list
        total_recipes.append(new_recipe)

        # isolate the ingredients list from the recipe
        new_ingredients = {new_name + ' ingredients': recipe_data['recipe']['ingredientLines']}

        # add isolated ingredients to list of ingredients
        just_ingredients.append(new_ingredients)

        # create new document
        doc = Document()
        # add heading
        doc.add_heading("Saved Recipes")

        # add each recipe to document
        for recipe_info in modified_data[0]:
            for result_number, recipe_details in recipe_info.items():
                title_paragraph = doc.add_paragraph()
                runner = title_paragraph.add_run(f"Recipe Title: {recipe_details['recipe']['label']}")
                runner.bold = True
                doc.add_paragraph(f"URL: {recipe_details['recipe']['url']}")

                # add ingredients as bulleted list
                doc.add_paragraph(f"Ingredients: ")
                for each_ingredient in recipe_details['recipe']['ingredientLines']:
                    ingredient_paragraph = doc.add_paragraph(f"{each_ingredient}")
                    ingredient_paragraph.style = 'List Bullet'

                doc.add_paragraph("\n")

        response_filename = 'Recipes.docx'
        doc.save(response_filename)


def exit_app():
    if messagebox.askokcancel("Quit", "Do you want to quit?"):
        root.destroy()


def main():
    global output_text
    global excluded_ingredients_entry
    global ingredients_entry
    global num_recipes_entry
    global saved_recipes_entry
    global root
    """
    ----------------------------------------------------------------------------
    Main Program
    ----------------------------------------------------------------------------
    """

    root = Tk()
    root.title("Meal Planner")

    bg_color = root.cget("bg")

    header_text = Text(root, height=1, width=len("Meal Planner App"), font=('Times New Roman', 48, 'bold'), bg=bg_color)
    header_text.insert(INSERT, "Meal Planner App!")
    header_text.tag_configure("center", justify=CENTER)
    header_text.tag_add("center", 1.0, "end")
    header_text.pack()

    instructions_label = Label(root, text="Enter ingredients to include (comma-separated):")
    instructions_label.pack()

    ingredients_entry = Entry(root, width=50)
    ingredients_entry.pack()

    excluded_ingredients_label = Label(root, text="Enter ingredients to exclude (comma-separated):")
    excluded_ingredients_label.pack()

    excluded_ingredients_entry = Entry(root, width=50)
    excluded_ingredients_entry.pack()

    num_recipes_label = Label(root, text="Enter the number of recipes (1 to 20):")
    num_recipes_label.pack()

    num_recipes_entry = Entry(root)
    num_recipes_entry.pack()

    button_frame = Frame(root)
    button_frame.pack()

    search_button = Button(button_frame, text="Search Recipes", command=search_recipes)
    search_button.pack(side=LEFT)

    browse_button = Button(button_frame, text="Browse Recipes", command=browse_recipes)
    browse_button.pack(side=RIGHT)

    saved_recipes_label = Label(root, text="Enter recipe numbers to save (Ex: 1, 2, 4)")
    saved_recipes_label.pack()

    saved_recipes_entry = Entry(root)
    saved_recipes_entry.pack()

    save_recipe_button = Button(text="Save Recipes", command=save_recipes)
    save_recipe_button.pack()

    export_recipe_to_word_button = Button(text="Export Saved Recipes to Word", command=create_recipe_document)
    export_recipe_to_word_button.pack()

    exit_button = Button(root, text="Exit", command=exit_app)
    exit_button.pack()

    output_text = scrolledtext.ScrolledText(root, width=70, height=50)
    output_text.pack()

    root.mainloop()


if __name__ == "__main__":
    main()
